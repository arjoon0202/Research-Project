from docx import Document
import copy

doc = Document('DISSERTATION_HUMANISED.docx')

# Each tuple: (para_index, old_fragment, new_fragment)
replacements = [
    # Para 122: single dash elaboration -> comma
    (122, ') — most obviously', '), most obviously'),
    # Para 123: paired parenthetical -> commas
    (123, 'developments — including', 'developments, including'),
    (123, 'regions — show', 'regions, show'),
    # Para 131: single dash consequence -> comma
    (131, 'settings — a practice', 'settings, a practice'),
    # Para 133: single dash -> comma
    (133, 'limitations — making', 'limitations, making'),
    # Para 134: single dash -> comma
    (134, 'cultures — a concern', 'cultures, a concern'),
    # Para 178: single dash -> comma
    (178, 'attitudes — including', 'attitudes, including'),
    # Para 182: single dash specifying -> comma
    (182, 'factors — both', 'factors, both'),
    # Para 184: single dash -> colon
    (184, 'inconsistent — affirming', 'inconsistent: affirming'),
    # Para 199: paired -> commas
    (199, 'inclusion — as opposed to mere physical integration — requires',
     'inclusion, as distinct from mere physical placement, requires'),
    # Para 201: single dash -> comma
    (201, 'effects — pointing', 'effects, pointing'),
    # Para 211: single dash -> comma
    (211, 'literature — particularly', 'literature, particularly'),
    # Para 221: single dash introducing list -> colon
    (221, "process — familiarisation", "process (familiarisation"),
    # Need to close the parenthetical for Para 221
    (221, "synthesis — provides", "synthesis) provides"),
    # Para 230: paired -> commas
    (230, 'documents — particularly', 'documents, particularly'),
    (230, 'Policy — through', 'Policy, through'),
    # Para 253: paired -> parentheses
    (253, 'concepts — reasonable', 'concepts (reasonable'),
    (253, 'Learning — it', 'Learning) it'),
    # Para 259: two pairs -> parentheses
    (259, 'practice — teacher preparation, specialist provision, and some regional coordination — show',
     'practice (teacher preparation, specialist provision, and some regional coordination) show'),
    (259, 'authority — systematic school-level coordination, multi-sector collaboration, and monitoring — show',
     'authority (systematic school-level coordination, multi-sector collaboration, and monitoring) show'),
    # Para 262: single -> comma
    (262, 'enacted — in some', 'enacted, in some'),
    # Para 276: single -> comma
    (276, '2004 — four', '2004, four'),
    # Para 283: paired -> colon + restructure
    (283, 'factors — the unadopted policy status, the gap between policy language and operational guidance, the limitations of professional development, the absence of monitoring systems, and governance incoherence — interact',
     'factors (the unadopted policy status, the gap between policy language and operational guidance, the limitations of professional development, the absence of monitoring systems, and governance incoherence) interact'),
    # Para 287: single -> comma
    (287, 'emerge — particularly', 'emerge, particularly'),
    # Para 291: paired -> commas
    (291, 'effect — in some cases more substantially than the policy\'s unadopted status would predict — but',
     'effect, in some cases more substantially than the policy\'s unadopted status would predict, but'),
    # Para 302: single -> comma
    (302, 'emerge — most', 'emerge, most'),
    # Para 305: paired -> parentheses
    (305, 'progress — 51 SEND-qualified graduates, 30 dedicated SEND spaces, and the formal CPCE programme launch — exceeds',
     'progress (51 SEND-qualified graduates, 30 dedicated SEND spaces, and the formal CPCE programme launch) exceeds'),
    # Para 310: paired -> commas
    (310, 'No. 4 — particularly reasonable accommodation and Universal Design for Learning — and',
     'No. 4, particularly reasonable accommodation and Universal Design for Learning, and'),
    # Para 325: single -> period + restructure
    (325, 'first step — not because', 'first step. This is not because'),
]

for para_idx, old, new in replacements:
    p = doc.paragraphs[para_idx]
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            break
    else:
        # The em dash might span runs; try full paragraph text replacement
        full = p.text
        if old in full:
            # Replace across runs by rebuilding
            combined = ''
            run_map = []
            for run in p.runs:
                for ch in run.text:
                    run_map.append((run, len(combined)))
                    combined += ch
            idx = combined.find(old)
            if idx >= 0:
                # Replace in combined text
                new_combined = combined[:idx] + new + combined[idx+len(old):]
                # Redistribute text across runs
                pos = 0
                for run in p.runs:
                    rlen = len(run.text)
                    run.text = new_combined[pos:pos+rlen]
                    pos += rlen
                # Handle length difference
                if len(new_combined) > pos:
                    p.runs[-1].text += new_combined[pos:]
                elif len(new_combined) < pos:
                    # Trim from last runs
                    excess = pos - len(new_combined)
                    for run in reversed(p.runs):
                        if excess <= 0:
                            break
                        if len(run.text) <= excess:
                            excess -= len(run.text)
                            run.text = ''
                        else:
                            run.text = run.text[:len(run.text)-excess]
                            excess = 0

doc.save('DISSERTATION_HUMANISED.docx')

# Verify
doc2 = Document('DISSERTATION_HUMANISED.docx')
count = sum(p.text.count('\u2014') for p in doc2.paragraphs)
print(f'Em dashes remaining: {count}')
if count > 0:
    for i, p in enumerate(doc2.paragraphs):
        if '\u2014' in p.text:
            print(f'  Para {i}: {p.text[:100]}...')
